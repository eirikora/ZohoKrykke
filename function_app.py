import azure.functions as func
import logging
import re
import os
import io
from datetime import datetime
import openai
import tempfile
import json
import base64
import mammoth
from bs4 import BeautifulSoup
import quopri
from docx import Document
from encryption import encrypt_word, decrypt_word
from dotenv import load_dotenv

load_dotenv()

app = func.FunctionApp(http_auth_level=func.AuthLevel.ANONYMOUS)

@app.route(route="http_test")
def http_test(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Python HTTP trigger function processed a request.')

    name = req.params.get('name')
    if not name:
        try:
            req_body = req.get_json()
        except ValueError:
            pass
        else:
            name = req_body.get('name')

    if name:
        encrypted_name = encrypt_word(name)
        decrypted_name = decrypt_word(encrypted_name)
        return func.HttpResponse(f"Hello, {decrypted_name}. Your name encrypted was {encrypted_name}. This HTTP triggered function executed successfully.")
    else:
        return func.HttpResponse(
             "This HTTP triggered function executed successfully. Pass a name in the query string or in the request body for a personalized response.",
             status_code=200
        )
    
@app.route(route="EncryptWord", auth_level=func.AuthLevel.ANONYMOUS)
def encrypt_word_function(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('EncryptWord trigger function processed a request.')

    # 1) Hent parameteren "secretword" fra query eller body
    secretword = req.params.get('secretword')
    if not secretword:
        try:
            req_body = req.get_json()
        except ValueError:
            req_body = {}
        secretword = req_body.get('secretword')

    if not secretword:
        return func.HttpResponse(
            "Please pass a 'secretword' either in querystring ?secretword=XXX or in JSON body {\"secretword\":\"XXX\"}.",
            status_code=400
        )

    # 2) Kall din funksjon for å kryptere
    result = encrypt_word(secretword)

    # 3) Returner kryptert streng til klienten
    return func.HttpResponse(result, status_code=200)

def decode_mime_string(s):
    if '=' in s:
        start = s.find('=')
        prefix = s[:start]
        #print(prefix)
        mimestring = s[start:]
        #print(mimestring)
        mime_elements = mimestring.split('=_=')
        #print(mime_elements)
        decoded_text = ''
        for submime in mime_elements:
            charset, encoded_text = submime.strip('=_').strip('_=').split('_Q_', 1)
            #decoded_text += quopri.decodestring(encoded_text.replace('_', '=')).decode(charset).replace('=', ' ')
            decoded_text += quopri.decodestring(encoded_text).decode(charset)
        return prefix + decoded_text
    else:
        return s


@app.route(route="MimeDecode", auth_level=func.AuthLevel.ANONYMOUS)
def MimeDecode(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('MimeDecode trigger function processed a request.')

    mime_string = req.params.get('mime_string')
    if not mime_string:
        try:
            req_body = req.get_json()
        except ValueError:
            mime_string = str(req.get_body(), 'utf-8')
        else:
            mime_string = req_body.get('mime_string')

    if mime_string:
        decoded_string = decode_mime_string(mime_string) # Now also handles prefixed mime-strings
        return func.HttpResponse(decoded_string, mimetype="text/plain;charset=UTF-8", status_code=200)
    else:
        return func.HttpResponse(
             "Please pass a mime_string on the query string or in the request body",
             status_code=400
        )
    
def allowed_file(filename):
    logging.info("Checking "+ filename)
    ALLOWED_EXTENSIONS = {'doc', 'docx'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_header_footer(docx_path):
    """
    Load a .docx file and extract the header and footer from the first page only.

    Args:
        docx_path (str): Path to the .docx file.

    Returns:
        str: Header and footer text from the first page, separated by newlines.
    """
    # Load the document
    doc = Document(docx_path)

    # Get the first section
    section = doc.sections[0]

    # If the document uses a different first page header/footer, use those
    if section.different_first_page_header_footer:
        header = section.first_page_header
        footer = section.first_page_footer
    else:
        header = section.header
        footer = section.footer

    # Extract text
    header_text = [p.text for p in header.paragraphs if p.text]
    footer_text = [p.text for p in footer.paragraphs if p.text]

    return "\n".join(header_text + footer_text)

@app.route(route="Word2Text", auth_level=func.AuthLevel.ANONYMOUS)
def Word2Text(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Word2Text trigger function processed a request.')

    """ # DEBUGGING
    # Grab method, URL, and headers
    method = req.method
    url    = req.url
    # Convert headers to a dict so you can log them as text
    headers_dict = dict(req.headers)

    # Get the body as bytes and decode to text with 'replace' to avoid errors if it's binary
    body_bytes = req.get_body()
    body_str   = body_bytes.decode("utf-8", errors="replace")

    # Build a "flattened" string: method, URL, headers, and the first 1000 chars of the body
    flattened_request_str = (
        f"\n--- REQUEST DEBUG ---\n"
        f"Method: {method}\n"
        f"URL: {url}\n"
        f"Headers: {headers_dict}\n"
        f"Body (first 1000 chars): {body_str[:1000]}\n"
        f"----------------------\n"
    )

    # Log it
    logging.info(flattened_request_str)

    # DEBUGGING """

    file_content = b''
    file_name = ''

    for field in req.files.keys():
        logging.info(f"DEBUG: Found req.files key: {field}")

    if 'content' in req.files:
        logging.info('Getting file from attached files...') # Typical for Zoho
        file = req.files['content']
        file_content = file.read()
        file_name = file.filename
    else:
        try:
            logging.info('Getting file from the raw body...') # Typical for Power Automate
            file_content = req.get_body()

            #DEBUG
            # Dekode til tekst, bytt ut evt. ugyldige tegn
            #body_str = req_body.decode('utf-8', errors='replace')
            # Logg en begrenset del av body (f.eks. 400 tegn)
            #logging.info(f"First 400 characters of request body:\n{body_str[:400]}")

            #json_data = json.loads(req_body, strict=False)
            #encoded_content = json_data["body"]["$content"]
            # Decode the content
            #file_content = base64.b64decode(encoded_content)
            file_name = "powerapp.docx"
        except ValueError:
            logging.info('ERROR: No file part in the request.')
            #file_content = req.get_body()
            return func.HttpResponse('ERROR: No file part in the request', mimetype="text/plain;charset=UTF-8", status_code=400)
        #return 'No file part in the request', 400
        #return func.HttpResponse(str(req_body[0:100]), mimetype="text/plain;charset=UTF-8", status_code=400)

    # Expect a Word filename to attempt convert
    if allowed_file(file_name):
        logging.info('Converting file ' + file_name)
        # Write the binary data to a temp Word file
        with tempfile.NamedTemporaryFile(mode='w+b', delete=False, suffix=".docx") as f:
            f.write(file_content)
            temp_filename = f.name
        
        # Get the header and footer from the file
        headerfooter = extract_header_footer(temp_filename)

        # Read and convert temp Word file into HTML for analysis
        with open(temp_filename, "rb") as docx_file:
            try:
                result = mammoth.convert_to_html(docx_file)
                html = result.value
            except Exception as e:
                html = "<p>Word2Text ERROR: Document was not a Microsoft Word document in .docx format (Zip file) that could be analyzed.</p>\n"
                # hex_representation = file_content[0:200].hex()
                # hex_representation_spaced = ' '.join(hex_representation[i:i+2] for i in range(0, len(hex_representation), 2))
                # html = html + "<p>" + hex_representation_spaced + "</p>\n"
                # html = html + "<p>" + str(req_body)[2:40] + "</p>\n"
        # Delete temp file
        try:
            os.remove(temp_filename)
        except OSError:
            pass
        
        # Create a BeautifulSoup object
        soup = BeautifulSoup(html, 'html.parser')

        plain_text = soup.get_text(separator='\n')
        plain_text = "HEADER_FOOTER:\n" + headerfooter + '\nFRONT_PAGE:\n' + plain_text
        plain_text = re.sub('\t',' ',plain_text)
        # Return the resultset
        return func.HttpResponse(plain_text, mimetype="text/plain;charset=UTF-8", status_code=200)
    else:
        logging.info('ERROR: Did not find a Word document to convert in request body!')
        return func.HttpResponse(
             "ERROR: Could not find a file in doc/docx format that we can convert. File received was: " + file_name,
             status_code=400
        )
    

@app.route(route="MakeVectorstore", auth_level=func.AuthLevel.ANONYMOUS)
def MakeVectorstore(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('MakeVectorstore trigger function processed a request.')

    encrypted_openai_key = req.params.get('openai_key')
    if not encrypted_openai_key:
        try:
            req_body = req.get_json()
        except ValueError:
            pass
        else:
            encrypted_openai_key = req_body.get('openai_key')
    openai_key = decrypt_word(encrypted_openai_key)

    vstore_name = req.params.get('vstore_name')
    if not vstore_name:
        try:
            req_body = req.get_json()
        except ValueError:
            pass
        else:
            vstore_name = req_body.get('vstore_name')

    if openai_key and vstore_name:
        logging.info("Connecting to OpenAI with supplied key.")
        try:
            OpenAIclient = openai.OpenAI(api_key=openai_key)
        except Exception as e:
            return func.HttpResponse( f"Failed to connect to OpenAI with supplied key: {e}",
             status_code=400
            )
        logging.info(f'Creating a vector store. {vstore_name}')
        # Get the current date and time
        current_datetime = datetime.now()
        formatted_datetime = current_datetime.strftime("%Y-%m-%d %H:%M:%S")
        try:
            vector_store = OpenAIclient.beta.vector_stores.create(
                name=f"{vstore_name} {formatted_datetime}"
            )
        except Exception as e:
            # Try to extract the error message from the response if available
            try:
                error_detail = e.response.json().get('error', {}).get('message', '')
            except Exception:
                error_detail = ''
            # Fallback to the full exception string if no nested error message is found
            if not error_detail:
                error_detail = str(e)

            return func.HttpResponse( f"Failed to create vector store with message: {error_detail}",
             status_code=400
            )
        logging.info(f"Created vector store with ID:{vector_store.id}")
        return_value = {
            "status": vector_store.status,
            "id": vector_store.id,
            "name": vector_store.name
        }
        return func.HttpResponse(json.dumps(return_value), mimetype="application/json", status_code=200)
    else:
        return func.HttpResponse(
             "Please pass a valid openai_key and vstore_name in the request parameters or body",
             status_code=400
        )
    

@app.route(route="UpsertTextdocVectorstore", auth_level=func.AuthLevel.ANONYMOUS)
def UpsertTextdocVectorstore(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('UpsertTextdocVectorstore trigger function processed a request.')
    # logging.info(f'REQUEST PARAMS:{req.params}')
    # logging.info(f'REQUEST FILES:{req.files.keys()}')
    
    """ # DEBUGGING
    # Grab method, URL, and headers
    method = req.method
    url    = req.url
    # Convert headers to a dict so you can log them as text
    headers_dict = dict(req.headers)

    # Get the body as bytes and decode to text with 'replace' to avoid errors if it's binary
    body_bytes = req.get_body()
    body_str   = body_bytes.decode("utf-8", errors="replace")

    # Build a "flattened" string: method, URL, headers, and the first 1000 chars of the body
    flattened_request_str = (
        f"\n--- REQUEST DEBUG ---\n"
        f"Method: {method}\n"
        f"URL: {url}\n"
        f"Headers: {headers_dict}\n"
        f"Body (first 1000 chars): {body_str[:1000]}\n"
        f"----------------------\n"
    )

    # Log it
    logging.info(flattened_request_str)

    # DEBUGGING """

    encrypted_openai_key = req.params.get('openai_key')
    if not encrypted_openai_key:
        try:
            req_body = req.get_json().get('body')
            encrypted_openai_key = req_body.get('openai_key')
        except ValueError:
            pass
            
    # logging.info(f'GOT openai key:{encrypted_openai_key}')
    openai_key = decrypt_word(encrypted_openai_key)

    vstore_id = req.params.get('vstore_id')
    if not vstore_id:
        try:
            req_body = req.get_json().get('body')
            vstore_id = req_body.get('vstore_id')
        except ValueError:
            pass
    
    file_name = req.params.get('file_name')
    if not file_name:
        try:
            req_body = req.get_json().get('body')
            file_name = req_body.get('file_name')
        except ValueError:
            pass
    # file_name extension MUST be lowercase for OpenAI
    name, ext = os.path.splitext(file_name)
    file_name = name + ext.lower()

    # logging.info(f'GOT file_name:{file_name}')
    file_length = 0
    file_content = req.params.get('file_content')
    if 'content' in req.files:
        logging.info('We have a request with files attached.')
        uploaded_file = req.files['content']
        file_content = uploaded_file.read()  # This should be bytes
        file_name = uploaded_file.filename
        if file_content is None:
            file_length = 0
        else:  
            file_length = len(file_content)
        logging.info(f"Read {file_length} bytes from request content/attachment '{file_name}'.")
    if file_content is None:
        file_length = 0
    else:  
        file_length = len(file_content)
    if file_length == 0:
        try:
            logging.info('Getting the body...')
            req_body = req.get_body()
            req_body_str = req_body.decode('utf-8')
            #logging.info(f'Looking at:{req_body_str}')
            if req_body_str.startswith('{'):
                logging.info('Getting the json body...')
                json_data = json.loads(req_body, strict=False)
                logging.info('Getting the json content field...')
                file_content = json_data["body"]["file_content"]
                # Decode the content
                # logging.info('Decoding the json body...')
                # file_content = base64.b64decode(encoded_content)
            else:
                logging.info("NO JSON HERE")
                file_content = req_body_str
            
        except ValueError:
            logging.info('ERROR: No file part in the request.')
            #file_content = req.get_body()
            return func.HttpResponse('ERROR: No file part in the request', mimetype="text/plain;charset=UTF-8", status_code=400)
        #return 'No file part in the request', 400
        #return func.HttpResponse(str(req_body[0:100]), mimetype="text/plain;charset=UTF-8", status_code=400)
    
    #logging.info(f"openai_key:{openai_key}")
    #logging.info(f"vstore_id:{vstore_id}")
    #logging.info(f"file_name:{file_name}")

    if openai_key and vstore_id and file_name:
        logging.info("Connecting to OpenAI with supplied key.")
        try:
            OpenAIclient = openai.OpenAI(api_key=openai_key)
        except Exception as e:
            return func.HttpResponse( f"Failed to connect to OpenAI with supplied key: {e}",
             status_code=400
            )
        logging.info(f'Connecting to vector store. {vstore_id}')
        try:
            vector_store = OpenAIclient.beta.vector_stores.retrieve(
                vector_store_id=vstore_id
            )
        except Exception as e:
            # Try to extract the error message from the response if available
            try:
                error_detail = e.response.json().get('error', {}).get('message', '')
            except Exception:
                error_detail = ''
            # Fallback to the full exception string if no nested error message is found
            if not error_detail:
                error_detail = str(e)

            return func.HttpResponse( f"Failed to identify vector store {vstore_id} with message: {error_detail}",
             status_code=400
            )
        
        logging.info(f"Checking if the file already exists and removing it if so.")
        allfiles = OpenAIclient.files.list()
        for openaifile in allfiles.data:
            if openaifile.filename == file_name:
                logging.info(f" - Found old file {file_name} with id {openaifile.id}, so removing it from Vectorstore.")
                # Remove it from the vector store
                try:
                    deleted_vector_store_file = OpenAIclient.beta.vector_stores.files.delete(
                        vector_store_id=vector_store.id,
                        file_id=openaifile.id
                    )
                except Exception:
                    pass
                # Delete the old uploaded file
                try:
                    deleted_file = OpenAIclient.files.delete(openaifile.id)
                except Exception:
                    pass

        logging.info(f"Adding new file to vector store with ID:{vector_store.id}")

        logging.info(f"NEW FILE NAME: {file_name}")
        logging.info(f"FILE CONTENT:{file_content[:80]}")

        # Opprett et fil-lignende objekt i minnet
        if isinstance(file_content, bytes):
            # If it's already bytes, just wrap it
            file_object = io.BytesIO(file_content)
        else:
            # Otherwise assume it's a string and encode to bytes
            file_object = io.BytesIO(file_content.encode('utf-8'))
        file_object.name = file_name  # Sett filnavnet (OpenAI only accepts lowercase name)

        logging.info(f"Uploading new file: {file_name}")

        # Last filen inn til OpenAI og få file_id
        OpenAIfile = OpenAIclient.files.create(
            file=file_object,
            purpose='assistants'
        )

        logging.info(f"Linking this file to vector store: {file_name}")
        # Last filen videre inn i Vectorstore
        vector_store_file = OpenAIclient.beta.vector_stores.files.create(
            vector_store_id=vector_store.id,
            file_id=OpenAIfile.id
        )

        logging.info(f"COMPLETED Linking this file to vector store: {file_name} and got status {vector_store_file.status}")
        return_value = {
            "status": vector_store_file.status,
            "vectorstore_id": vector_store.id,
            "vectorstore_name": vector_store.name,
            "file_name": file_name,
            "file_id": OpenAIfile.id
        }
        return func.HttpResponse(json.dumps(return_value), mimetype="application/json", status_code=200)
    else:
        return func.HttpResponse(
             "Please pass a valid openai_key, vstore_id and file_name in the request parameters or body",
             status_code=400
        )

@app.route(route="UpdateAssistantVectorstore", auth_level=func.AuthLevel.ANONYMOUS)
def UpdateAssistantVectorstore(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('UpdateAssistantVectorstore trigger function processed a request.')
    # logging.info(f'REQUEST PARAMS:{req.params}')

    encrypted_openai_key = req.params.get('openai_key')
    if not encrypted_openai_key:
        try:
            req_body = req.get_json()
        except ValueError:
            pass
        else:
            encrypted_openai_key = req_body.get('openai_key')
    openai_key = decrypt_word(encrypted_openai_key)

    assistant_id = req.params.get('assistant_id')
    if not assistant_id:
        try:
            req_body = req.get_json()
        except ValueError:
            pass
        else:
            assistant_id = req_body.get('assistant_id')

    vstore_id = req.params.get('vstore_id')
    if not vstore_id:
        try:
            req_body = req.get_json()
        except ValueError:
            pass
        else:
            vstore_id = req_body.get('vstore_id')

    if openai_key and assistant_id and vstore_id:

        logging.info("Connecting to OpenAI with supplied key.")
        try:
            OpenAIclient = openai.OpenAI(api_key=openai_key)
        except Exception as e:
            return func.HttpResponse( f"Failed to connect to OpenAI with supplied key: {e}",
             status_code=400
            )
        
        # Oppdater assistenten med den nye Vector Store
        try:
            updated_assistant = OpenAIclient.beta.assistants.update(
                assistant_id=assistant_id,
                tool_resources={
                    "file_search": {
                        "vector_store_ids": [vstore_id]
                    }
                }
            )
        except Exception as e:
            return func.HttpResponse( f"Failed to replace vectorstore to id {vstore_id} for assistant {assistant_id} with supplied key: {e}",
                status_code=400
            )
        
        return_value = {
            "operation": f"Assistant {updated_assistant.id} is now using vectorstore {vstore_id}.",
            "status": "success",
            "assistant_id": updated_assistant.id,
            "vectorstore_id": vstore_id
        }
        return func.HttpResponse(json.dumps(return_value), mimetype="application/json", status_code=200)

    else:
        return func.HttpResponse(
             "Please pass a valid openai_key, assistant_id and vstore_name in the request parameters.",
             status_code=400
        )