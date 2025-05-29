import sys
from docx import Document

def extract_header_footer_first_page(docx_path):
    """
    Load a .docx file and extract the header and footer from the first page only.

    Args:
        docx_path (str): Path to the .docx file.

    Returns:
        str: Header and footer text from the first page, separated by newlines.
    """
    # Load the document
    doc = Document(docx_path)

    # Get the first section
    section = doc.sections[0]

    # If the document uses a different first page header/footer, use those
    if section.different_first_page_header_footer:
        header = section.first_page_header
        footer = section.first_page_footer
    else:
        header = section.header
        footer = section.footer

    # Extract text
    header_text = [p.text for p in header.paragraphs if p.text]
    footer_text = [p.text for p in footer.paragraphs if p.text]

    return "\n".join(header_text + footer_text)


def main():
    if len(sys.argv) != 2:
        print(f"Usage: python {sys.argv[0]} <path-to-docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    try:
        result = extract_header_footer_first_page(docx_path)
        print("--- First Page Header & Footer ---")
        print(result if result else "(No header or footer found on first page)")
    except Exception as e:
        print(f"Error reading '{docx_path}': {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
